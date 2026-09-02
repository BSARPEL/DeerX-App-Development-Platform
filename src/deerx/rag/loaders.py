"""Dokuman yukleyicileri: dosyayi duz metne ve meta veriye cevirir."""

from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ..errors import IngestError
from ..logging import get_logger

log = get_logger("rag.loaders")

CODE_SUFFIXES = {
    ".py", ".ts", ".tsx", ".js", ".jsx", ".go", ".rs", ".java", ".cs", ".kt",
    ".rb", ".php", ".c", ".h", ".cpp", ".hpp", ".swift", ".scala", ".sh", ".sql",
}
DATA_SUFFIXES = {".json", ".yaml", ".yml", ".toml", ".ini", ".csv", ".xml"}
DOC_SUFFIXES = {".md", ".txt", ".rst", ".markdown", ".adoc"}
BINARY_DOC_SUFFIXES = {".pdf", ".docx"}
WEB_SUFFIXES = {".html", ".htm"}

SUPPORTED_SUFFIXES = (
    CODE_SUFFIXES | DATA_SUFFIXES | DOC_SUFFIXES | BINARY_DOC_SUFFIXES | WEB_SUFFIXES
)


@dataclass(slots=True)
class LoadedDoc:
    """Bir kaynagin yuklenmis, henuz parcalanmamis hali."""

    source: str
    title: str
    kind: str  # doc | code | data | web
    text: str
    path: Path | None = None
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def sha256(self) -> str:
        return hashlib.sha256(self.text.encode("utf-8", "replace")).hexdigest()


def classify(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in CODE_SUFFIXES:
        return "code"
    if suffix in DATA_SUFFIXES:
        return "data"
    if suffix in WEB_SUFFIXES:
        return "web"
    return "doc"


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_SUFFIXES


# ---------------------------------------------------------------------- #
# Bicime ozgu cikaricilar
# ---------------------------------------------------------------------- #
def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1254", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise IngestError(f"{path} metin olarak cozulemedi (kodlama bilinmiyor).")


def _read_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise IngestError("PDF okumak icin `pypdf` gerekli: uv add pypdf") from exc

    pages: list[str] = []
    try:
        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                # Sayfa isaretcisi, alintilarda kaynak gostermeyi mumkun kilar.
                pages.append(f"\n\n[sayfa {index}]\n{text}")
    except Exception as exc:  # noqa: BLE001 - pypdf birden cok hata tipi firlatir
        # Bozuk tek bir dosya butun indeksleme fazini dusurmemeli.
        raise IngestError(f"{path.name} gecerli bir PDF gibi okunamadi: {exc}") from exc

    if not any(p.strip() for p in pages):
        raise IngestError(
            f"{path.name} icinden metin cikarilamadi (taranmis PDF olabilir; OCR gerekir)."
        )
    return "".join(pages), {"pages": len(reader.pages)}


def _read_docx(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover
        raise IngestError("DOCX okumak icin `python-docx` gerekli.") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 - python-docx kendi hata tiplerini kullanir
        raise IngestError(f"{path.name} gecerli bir DOCX degil: {exc}") from exc

    parts: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            # Word basliklarini markdown basligina cevir; parcalayici bunu kullanir.
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            parts.append(f"\n[tablo {table_index}]\n" + "\n".join(rows))
    return "\n\n".join(parts), {"paragraphs": len(document.paragraphs)}


def _read_html(path_or_text: Path | str) -> str:
    from bs4 import BeautifulSoup

    raw = path_or_text if isinstance(path_or_text, str) else _read_text(path_or_text)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "svg"]):
        tag.decompose()
    # Baslik hiyerarsisini markdown'a cevir ki parcalayici yapiyi gorebilsin.
    for level in range(1, 7):
        for tag in soup.find_all(f"h{level}"):
            tag.replace_with(f"\n\n{'#' * level} {tag.get_text(' ', strip=True)}\n\n")
    text = soup.get_text("\n", strip=True)
    return "\n".join(line for line in text.splitlines() if line.strip())


# ---------------------------------------------------------------------- #
# Genel giris noktalari
# ---------------------------------------------------------------------- #
def load_file(path: Path, *, max_bytes: int = 2_000_000) -> LoadedDoc:
    """Tek bir dosyayi yukler."""
    path = path.resolve()
    if not path.is_file():
        raise IngestError(f"Dosya bulunamadi: {path}")

    size = path.stat().st_size
    if size > max_bytes:
        raise IngestError(f"{path.name} cok buyuk ({size:,} bayt > {max_bytes:,}).")

    suffix = path.suffix.lower()
    meta: dict[str, Any] = {"bytes": size, "suffix": suffix}

    if suffix == ".pdf":
        text, extra = _read_pdf(path)
        meta.update(extra)
    elif suffix == ".docx":
        text, extra = _read_docx(path)
        meta.update(extra)
    elif suffix in WEB_SUFFIXES:
        text = _read_html(path)
    elif suffix in DATA_SUFFIXES and suffix == ".json":
        raw = _read_text(path)
        try:
            # JSON'u yeniden bicimlendirmek satir bazli alintilamayi kolaylastirir.
            text = json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            text = raw
    else:
        text = _read_text(path)

    if not text.strip():
        raise IngestError(f"{path.name} bos.")

    return LoadedDoc(
        source=str(path),
        title=path.name,
        kind=classify(path),
        text=text,
        path=path,
        meta=meta,
    )


def load_html_text(html: str, *, source: str, title: str | None = None) -> LoadedDoc:
    """Web'den cekilen ham HTML'i yuklenmis dokumana cevirir."""
    text = _read_html(html)
    if not text.strip():
        raise IngestError(f"{source} icinden metin cikarilamadi.")
    return LoadedDoc(
        source=source,
        title=title or source,
        kind="web",
        text=text,
        meta={"origin": "web"},
    )


def load_text(text: str, *, source: str, title: str, kind: str = "doc", **meta: Any) -> LoadedDoc:
    """Bellekteki metni dogrudan dokuman olarak sarar."""
    return LoadedDoc(source=source, title=title, kind=kind, text=text, meta=meta)


def iter_files(
    root: Path,
    include: list[str],
    exclude: list[str],
) -> list[Path]:
    """`include` glob'lariyla eslesen, `exclude` ile elenmemis dosyalari listeler."""
    import fnmatch

    root = root.resolve()
    seen: dict[Path, None] = {}
    for pattern in include:
        for candidate in root.glob(pattern):
            if candidate.is_file():
                seen.setdefault(candidate.resolve(), None)

    def blocked(path: Path) -> bool:
        rel = path.relative_to(root).as_posix() if path.is_relative_to(root) else path.as_posix()
        return any(fnmatch.fnmatch(rel, pat.lstrip("/")) or fnmatch.fnmatch(f"/{rel}", pat)
                   for pat in exclude)

    return sorted(p for p in seen if not blocked(p))
